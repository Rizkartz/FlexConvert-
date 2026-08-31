import os
import markdown
import html2text
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import PyPDF2

def convert_document(input_path: str, output_format: str, output_dir: str) -> str:
    filename = os.path.basename(input_path)
    base_name, ext = os.path.splitext(filename)
    ext = ext.lower().replace('.', '')
    output_filename = f"{base_name}.{output_format}"
    output_path = os.path.join(output_dir, output_filename)
    
    text_content = ""
    
    # Read input
    if ext == 'docx':
        doc = Document(input_path)
        text_content = "\n".join([para.text for para in doc.paragraphs])
    elif ext in ['txt', 'md', 'html']:
        with open(input_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
    elif ext == 'pdf':
        with open(input_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content += extracted + "\n"
    
    # Optional intermediate conversions
    if ext == 'md' and output_format == 'html':
        text_content = markdown.markdown(text_content)
    elif ext == 'html' and output_format in ['txt', 'md']:
        h = html2text.HTML2Text()
        h.ignore_links = False
        text_content = h.handle(text_content)
        
    # Write output
    if output_format == 'pdf':
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        y = height - 40
        for line in text_content.split('\n'):
            # simple text wrapping could be added here, keeping it basic for now
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line[:100]) # naive crop to avoid running off edge
            y -= 15
        c.save()
    elif output_format == 'docx':
        doc = Document()
        for line in text_content.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)
    elif output_format in ['txt', 'md', 'html']:
        if output_format == 'html' and ext != 'md' and ext != 'html':
            # Simple wrapper
            text_content = f"<html><body>\n" + "<br>\n".join(text_content.split('\n')) + "\n</body></html>"
            
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
            
    return output_path
